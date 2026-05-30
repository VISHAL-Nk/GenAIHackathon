"""
DOCX Score Sheet Generator Module
====================================
Generates professional DOCX score sheets for viva voce examinations.

Creates a polished, color-coded document with:
    - Institution-style header with examination title and date
    - Student information table with colored header cells
    - Question-wise scoring table with Bloom's level color coding
    - Score summary with bloom-wise breakdown and final grade
    - Professional footer with timestamp

Functions:
    generate_score_sheet(...) — Main function to create the DOCX file
"""

import os
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Color constants (hex without '#' prefix for DOCX XML)
# ---------------------------------------------------------------------------
NAVY_DARK = "1E3A5F"       # Dark navy for header cells
NAVY_MEDIUM = "2C5282"     # Medium navy for sub-headers
GRAY_LIGHT = "F7FAFC"      # Light gray for value cells
GRAY_BORDER = "CBD5E0"     # Gray for table borders
WHITE = "FFFFFF"

# Bloom level colors (background)
BLOOM_COLORS = {
    "BL2": "DBEAFE",  # Light blue
    "BL3": "D1FAE5",  # Light green
    "BL4": "FEF3C7",  # Light yellow
    "BL5": "FEE2E2",  # Light red
}

# Score colors (background)
SCORE_COLORS_HIGH = "C8F7C5"   # Light green for score 4-5
SCORE_COLORS_MED = "FFF3CD"    # Light yellow for score 2-3
SCORE_COLORS_LOW = "FFDEDE"    # Light red for score 0-1

# Grade scale
GRADE_SCALE = [
    (90, "O (Outstanding)"),
    (80, "A+ (Excellent)"),
    (70, "A (Very Good)"),
    (60, "B+ (Good)"),
    (50, "B (Average)"),
    (40, "C (Pass)"),
    (0, "F (Fail)"),
]


# ---------------------------------------------------------------------------
# Helper: Set cell background shading
# ---------------------------------------------------------------------------
def _set_cell_shading(cell, color_hex: str):
    """
    Set the background/shading color of a table cell.

    Uses low-level OOXML manipulation because python-docx doesn't expose
    cell shading directly. Creates a <w:shd> element inside the cell's
    table cell properties (<w:tcPr>).

    Parameters:
        cell: A python-docx table cell object.
        color_hex (str): Hex color string WITHOUT '#' prefix (e.g., '1E3A5F').
    """
    # Create the shading XML element: <w:shd w:fill="COLOR" w:val="clear"/>
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color_hex)
    shading_elm.set(qn("w:val"), "clear")
    # Append to the cell's table cell properties (tcPr)
    # get_or_add_tcPr() ensures the <w:tcPr> element exists
    cell._tc.get_or_add_tcPr().append(shading_elm)


# ---------------------------------------------------------------------------
# Helper: Set cell text with formatting
# ---------------------------------------------------------------------------
def _set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    font_size: int = 10,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    font_color: str = None,
    italic: bool = False,
):
    """
    Set text in a table cell with formatting options.

    Clears any existing paragraphs and adds a new one with the specified
    formatting. Uses Calibri font throughout for consistency.

    Parameters:
        cell: A python-docx table cell object.
        text (str): Text content to set.
        bold (bool): Whether to bold the text.
        font_size (int): Font size in points.
        alignment: Paragraph alignment (WD_ALIGN_PARAGRAPH enum).
        font_color (str): Hex color WITHOUT '#' (e.g., 'FFFFFF'). None for default.
        italic (bool): Whether to italicize the text.
    """
    # Clear existing content
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    # Reduce spacing for compact cells
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)

    run = paragraph.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic

    if font_color:
        # Parse hex color string into RGB components
        r = int(font_color[0:2], 16)
        g = int(font_color[2:4], 16)
        b = int(font_color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


# ---------------------------------------------------------------------------
# Helper: Add a second run to an existing cell (for hints in smaller text)
# ---------------------------------------------------------------------------
def _add_cell_text(
    cell,
    text: str,
    font_size: int = 8,
    italic: bool = True,
    font_color: str = "666666",
):
    """
    Append additional text to an existing cell as a new paragraph.

    Used for adding model answer hints below the justification text
    in a smaller, italic font.

    Parameters:
        cell: A python-docx table cell object.
        text (str): Text to append.
        font_size (int): Font size in points.
        italic (bool): Whether to italicize.
        font_color (str): Hex color WITHOUT '#'.
    """
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)
    run.font.italic = italic
    if font_color:
        r = int(font_color[0:2], 16)
        g = int(font_color[2:4], 16)
        b = int(font_color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


# ---------------------------------------------------------------------------
# Helper: Get Bloom level color
# ---------------------------------------------------------------------------
def _get_bloom_color_hex(bloom_level: str) -> str:
    """
    Return the hex background color for a Bloom's Taxonomy level.

    Parameters:
        bloom_level (str): 'BL2', 'BL3', 'BL4', or 'BL5'.

    Returns:
        str: Hex color string without '#' prefix.
    """
    return BLOOM_COLORS.get(bloom_level, "F3F4F6")


# ---------------------------------------------------------------------------
# Helper: Get score color
# ---------------------------------------------------------------------------
def _get_score_color_hex(score: int) -> str:
    """
    Return the hex background color for a score value.

    Parameters:
        score (int): Score from 0 to 5.

    Returns:
        str: Hex color string without '#' prefix.
    """
    if score >= 4:
        return SCORE_COLORS_HIGH
    elif score >= 2:
        return SCORE_COLORS_MED
    else:
        return SCORE_COLORS_LOW


# ---------------------------------------------------------------------------
# Helper: Calculate grade from percentage
# ---------------------------------------------------------------------------
def _calculate_grade(percentage: float) -> str:
    """
    Determine the letter grade from a percentage.

    Parameters:
        percentage (float): Score percentage (0-100).

    Returns:
        str: Grade string (e.g., 'A+ (Excellent)').
    """
    for threshold, grade in GRADE_SCALE:
        if percentage >= threshold:
            return grade
    return "F (Fail)"


# ---------------------------------------------------------------------------
# Helper: Set table borders
# ---------------------------------------------------------------------------
def _set_table_borders(table, color: str = "999999", size: int = 4):
    """
    Set borders for all cells in a table.

    Parameters:
        table: A python-docx table object.
        color (str): Border color hex without '#'.
        size (int): Border width in eighths of a point.
    """
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)

    tbl_pr.append(borders)


# ---------------------------------------------------------------------------
# Main function: Generate the score sheet
# ---------------------------------------------------------------------------
def generate_score_sheet(
    student_name: str,
    student_id: str,
    subject: str,
    department: str,
    report_title: str,
    faculty_name: str,
    questions_data: list,
    eval_results: list,
    output_path: str,
) -> str:
    """
    Generate a professional DOCX score sheet for the viva voce examination.

    Creates a formatted document with student info, all questions with scores,
    bloom-wise breakdown, and final grade summary.

    Parameters:
        student_name (str): Full name of the student.
        student_id (str): Student's registration/ID number.
        subject (str): Subject or course name.
        department (str): Student's department.
        report_title (str): Title of the submitted project report.
        faculty_name (str): Name of the supervising faculty.
        questions_data (list[dict]): List of question dicts with keys:
            q_number, bloom_level, bloom_label, question
        eval_results (list[dict]): List of evaluation dicts with keys:
            score, justification, correct_answer_hint
            (aligned by index with questions_data)
        output_path (str): File path to save the DOCX file.

    Returns:
        str: The output_path where the file was saved.
    """
    # --- Ensure output directory exists ---
    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)

    doc = Document()

    # --- Set default font ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # ======================================================================
    # HEADER SECTION
    # ======================================================================
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(4)
    title_run = title_para.add_run("VIVA VOCE EXAMINATION — SCORE SHEET")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = "Calibri"
    title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)  # Dark navy

    # Subtitle with date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_after = Pt(12)
    date_run = date_para.add_run(
        f"Date: {datetime.now().strftime('%d %B %Y')}"
    )
    date_run.font.size = Pt(11)
    date_run.font.name = "Calibri"
    date_run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)  # Gray

    # ======================================================================
    # STUDENT INFORMATION TABLE
    # ======================================================================
    info_heading = doc.add_paragraph()
    info_heading.paragraph_format.space_before = Pt(8)
    info_heading.paragraph_format.space_after = Pt(4)
    info_run = info_heading.add_run("STUDENT INFORMATION")
    info_run.font.size = Pt(12)
    info_run.font.bold = True
    info_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    info_table = doc.add_table(rows=4, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(info_table)

    # Row 1: Student Name | value | Student ID | value
    _set_cell_text(info_table.rows[0].cells[0], "Student Name", bold=True, font_size=10, font_color=WHITE)
    _set_cell_shading(info_table.rows[0].cells[0], NAVY_DARK)
    _set_cell_text(info_table.rows[0].cells[1], student_name, font_size=10)
    _set_cell_shading(info_table.rows[0].cells[1], GRAY_LIGHT)
    _set_cell_text(info_table.rows[0].cells[2], "Student ID", bold=True, font_size=10, font_color=WHITE)
    _set_cell_shading(info_table.rows[0].cells[2], NAVY_DARK)
    _set_cell_text(info_table.rows[0].cells[3], student_id, font_size=10)
    _set_cell_shading(info_table.rows[0].cells[3], GRAY_LIGHT)

    # Row 2: Subject | value | Department | value
    _set_cell_text(info_table.rows[1].cells[0], "Subject", bold=True, font_size=10, font_color=WHITE)
    _set_cell_shading(info_table.rows[1].cells[0], NAVY_DARK)
    _set_cell_text(info_table.rows[1].cells[1], subject, font_size=10)
    _set_cell_shading(info_table.rows[1].cells[1], GRAY_LIGHT)
    _set_cell_text(info_table.rows[1].cells[2], "Department", bold=True, font_size=10, font_color=WHITE)
    _set_cell_shading(info_table.rows[1].cells[2], NAVY_DARK)
    _set_cell_text(info_table.rows[1].cells[3], department, font_size=10)
    _set_cell_shading(info_table.rows[1].cells[3], GRAY_LIGHT)

    # Row 3: Report Title | value (merged across 3 cells)
    _set_cell_text(info_table.rows[2].cells[0], "Report Title", bold=True, font_size=10, font_color=WHITE)
    _set_cell_shading(info_table.rows[2].cells[0], NAVY_DARK)
    # Merge cells 1-3 for the report title
    merged_cell = info_table.rows[2].cells[1].merge(info_table.rows[2].cells[3])
    _set_cell_text(merged_cell, report_title, font_size=10)
    _set_cell_shading(merged_cell, GRAY_LIGHT)

    # Row 4: Faculty Name | value (merged across 3 cells)
    _set_cell_text(info_table.rows[3].cells[0], "Faculty Name", bold=True, font_size=10, font_color=WHITE)
    _set_cell_shading(info_table.rows[3].cells[0], NAVY_DARK)
    merged_cell = info_table.rows[3].cells[1].merge(info_table.rows[3].cells[3])
    _set_cell_text(merged_cell, faculty_name, font_size=10)
    _set_cell_shading(merged_cell, GRAY_LIGHT)

    # Add spacing after info table
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ======================================================================
    # QUESTION-WISE SCORING TABLE
    # ======================================================================
    q_heading = doc.add_paragraph()
    q_heading.paragraph_format.space_before = Pt(8)
    q_heading.paragraph_format.space_after = Pt(4)
    q_run = q_heading.add_run("QUESTION-WISE EVALUATION")
    q_run.font.size = Pt(12)
    q_run.font.bold = True
    q_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    # Create table: header + 10 question rows
    num_questions = len(questions_data)
    q_table = doc.add_table(rows=num_questions + 1, cols=6)
    q_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(q_table)

    # Header row
    headers = ["Q#", "Bloom Level", "Question", "Student Answer", "Score", "Justification"]
    for i, header in enumerate(headers):
        _set_cell_text(
            q_table.rows[0].cells[i], header,
            bold=True, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_color=WHITE
        )
        _set_cell_shading(q_table.rows[0].cells[i], NAVY_DARK)

    # Data rows
    for idx in range(num_questions):
        q = questions_data[idx]
        row = q_table.rows[idx + 1]

        # Get evaluation result (handle case where eval might not exist yet)
        if idx < len(eval_results) and eval_results[idx]:
            ev = eval_results[idx]
        else:
            ev = {"score": 0, "justification": "Not evaluated", "correct_answer_hint": "N/A"}

        # Get student answer from eval_results or default
        student_answer = ev.get("student_answer", "No answer provided")

        # Q# column
        _set_cell_text(row.cells[0], str(q.get("q_number", idx + 1)),
                       font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

        # Bloom Level column — color-coded
        bloom = q.get("bloom_level", "BL2")
        bloom_label = q.get("bloom_label", "Understand")
        _set_cell_text(row.cells[1], f"{bloom}\n({bloom_label})",
                       font_size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        _set_cell_shading(row.cells[1], _get_bloom_color_hex(bloom))

        # Question column
        _set_cell_text(row.cells[2], q.get("question", ""), font_size=9)

        # Student Answer column
        _set_cell_text(row.cells[3], student_answer, font_size=9)

        # Score column — color-coded
        score = ev.get("score", 0)
        _set_cell_text(row.cells[4], f"{score}/5",
                       font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        _set_cell_shading(row.cells[4], _get_score_color_hex(score))

        # Justification column — includes model answer hint
        _set_cell_text(row.cells[5], ev.get("justification", "N/A"), font_size=8)
        hint = ev.get("correct_answer_hint", "")
        if hint and hint != "N/A":
            _add_cell_text(row.cells[5], f"📝 Model Answer: {hint}",
                          font_size=7, italic=True, font_color="4A5568")

    # Set approximate column widths
    # (python-docx column widths are set per-cell in the first row)
    widths = [Cm(1.2), Cm(2.5), Cm(5.0), Cm(4.5), Cm(1.5), Cm(5.0)]
    for i, width in enumerate(widths):
        for row in q_table.rows:
            row.cells[i].width = width

    # Add spacing
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ======================================================================
    # SCORE SUMMARY TABLE
    # ======================================================================
    summary_heading = doc.add_paragraph()
    summary_heading.paragraph_format.space_before = Pt(8)
    summary_heading.paragraph_format.space_after = Pt(4)
    summary_run = summary_heading.add_run("SCORE SUMMARY")
    summary_run.font.size = Pt(12)
    summary_run.font.bold = True
    summary_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    # Calculate bloom-wise scores
    bloom_scores = {"BL2": 0, "BL3": 0, "BL4": 0, "BL5": 0}
    bloom_max = {"BL2": 15, "BL3": 15, "BL4": 10, "BL5": 10}  # 3×5, 3×5, 2×5, 2×5
    total_score = 0

    for idx, q in enumerate(questions_data):
        bl = q.get("bloom_level", "BL2")
        score = 0
        if idx < len(eval_results) and eval_results[idx]:
            score = eval_results[idx].get("score", 0)
        bloom_scores[bl] = bloom_scores.get(bl, 0) + score
        total_score += score

    max_score = 50
    percentage = round((total_score / max_score) * 100, 1) if max_score > 0 else 0
    grade = _calculate_grade(percentage)
    passed = percentage >= 40

    # Create summary table (8 rows × 2 cols)
    summary_table = doc.add_table(rows=9, cols=2)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(summary_table)

    summary_data = [
        ("BL2 — Understand (3 Questions)", f"{bloom_scores['BL2']} / {bloom_max['BL2']}"),
        ("BL3 — Apply (3 Questions)", f"{bloom_scores['BL3']} / {bloom_max['BL3']}"),
        ("BL4 — Analyze (2 Questions)", f"{bloom_scores['BL4']} / {bloom_max['BL4']}"),
        ("BL5 — Evaluate (2 Questions)", f"{bloom_scores['BL5']} / {bloom_max['BL5']}"),
        ("Total Score", f"{total_score} / {max_score}"),
        ("Percentage", f"{percentage}%"),
        ("Grade", grade),
        ("Result", "PASS ✓" if passed else "FAIL ✗"),
        ("Faculty Signature", ""),
    ]

    for i, (label, value) in enumerate(summary_data):
        _set_cell_text(summary_table.rows[i].cells[0], label,
                       bold=True, font_size=10, font_color=WHITE)
        _set_cell_shading(summary_table.rows[i].cells[0], NAVY_DARK)
        _set_cell_text(summary_table.rows[i].cells[1], value,
                       font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(summary_table.rows[i].cells[1], GRAY_LIGHT)

        # Color-code the result row
        if label == "Result":
            result_color = "C8F7C5" if passed else "FFDEDE"
            _set_cell_shading(summary_table.rows[i].cells[1], result_color)

    # Set column widths for summary table
    for row in summary_table.rows:
        row.cells[0].width = Cm(8)
        row.cells[1].width = Cm(6)

    # ======================================================================
    # FOOTER
    # ======================================================================
    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Separator line
    sep_para = doc.add_paragraph()
    sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep_para.add_run("─" * 60)
    sep_run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE0)

    # Footer text
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_before = Pt(4)
    footer_run = footer_para.add_run(
        "Generated by Automated Viva Voce Examiner"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    # Timestamp
    ts_para = doc.add_paragraph()
    ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts_run = ts_para.add_run(
        f"Generated on: {datetime.now().strftime('%d %B %Y at %I:%M %p')}"
    )
    ts_run.font.size = Pt(8)
    ts_run.font.italic = True
    ts_run.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

    # ======================================================================
    # SAVE DOCUMENT
    # ======================================================================
    doc.save(output_path)
    print(f"[DOCX] Score sheet saved to: {output_path}")

    return output_path


# ---------------------------------------------------------------------------
# Standalone smoke test
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    print("=" * 60)
    print("DOCX Generator — Smoke Test")
    print("=" * 60)

    # Create sample data
    sample_questions = [
        {"q_number": 1, "bloom_level": "BL2", "bloom_label": "Understand",
         "question": "Explain the main objective of your project and how it addresses the problem statement."},
        {"q_number": 2, "bloom_level": "BL2", "bloom_label": "Understand",
         "question": "Describe the key technologies used in your implementation."},
        {"q_number": 3, "bloom_level": "BL2", "bloom_label": "Understand",
         "question": "Summarize the literature survey findings relevant to your work."},
        {"q_number": 4, "bloom_level": "BL3", "bloom_label": "Apply",
         "question": "Demonstrate how you would modify your system to handle a different dataset."},
        {"q_number": 5, "bloom_level": "BL3", "bloom_label": "Apply",
         "question": "Apply the concepts from your project to solve a related real-world problem."},
        {"q_number": 6, "bloom_level": "BL3", "bloom_label": "Apply",
         "question": "Use your project methodology to propose a solution for scalability issues."},
        {"q_number": 7, "bloom_level": "BL4", "bloom_label": "Analyze",
         "question": "Compare your approach with at least two alternative methods discussed in literature."},
        {"q_number": 8, "bloom_level": "BL4", "bloom_label": "Analyze",
         "question": "Examine the performance bottlenecks in your system and their root causes."},
        {"q_number": 9, "bloom_level": "BL5", "bloom_label": "Evaluate",
         "question": "Justify why your chosen algorithm is the most suitable for this problem."},
        {"q_number": 10, "bloom_level": "BL5", "bloom_label": "Evaluate",
         "question": "Critically assess the limitations of your project and propose improvements."},
    ]

    sample_results = [
        {"score": 5, "justification": "Excellent explanation with proper terminology and deep understanding.",
         "correct_answer_hint": "The project aims to automate viva voce examinations using AI.",
         "student_answer": "The project automates viva voce question generation using LLMs and Bloom's Taxonomy."},
        {"score": 4, "justification": "Good coverage of technologies but missed mentioning the evaluation component.",
         "correct_answer_hint": "Key technologies include Streamlit, Groq API, PyPDF2, and python-docx.",
         "student_answer": "We used Streamlit for UI, Groq for LLM, and PyPDF2 for PDF parsing."},
        {"score": 3, "justification": "Main findings mentioned but lacks depth in analysis.",
         "correct_answer_hint": "Literature survey covers Bloom's Taxonomy applications in automated assessment.",
         "student_answer": "The literature shows that automated assessment is an active research area."},
        {"score": 4, "justification": "Good demonstration of adaptability with concrete steps.",
         "correct_answer_hint": "Modify the PDF extractor and prompt template for the new dataset format.",
         "student_answer": "I would update the text extraction pipeline and adjust the prompt templates."},
        {"score": 2, "justification": "Attempted but the application was too vague and generic.",
         "correct_answer_hint": "Apply the Bloom's Taxonomy mapping to create assessment tools for other subjects.",
         "student_answer": "It could be used in other areas of education."},
        {"score": 5, "justification": "Comprehensive solution with specific scalability measures.",
         "correct_answer_hint": "Use caching, async processing, and load balancing for scalability.",
         "student_answer": "I would implement caching for API responses, use async calls, and add load balancing."},
        {"score": 3, "justification": "Comparison made but only superficially — lacks quantitative analysis.",
         "correct_answer_hint": "Compare with rule-based and ML-based approaches on accuracy and coverage.",
         "student_answer": "Rule-based methods are simpler but less flexible than our LLM approach."},
        {"score": 1, "justification": "Barely addresses the question — no specific bottlenecks identified.",
         "correct_answer_hint": "Key bottlenecks include API latency, PDF parsing of scanned docs, and token limits.",
         "student_answer": "The system is sometimes slow."},
        {"score": 4, "justification": "Strong justification with relevant technical reasoning.",
         "correct_answer_hint": "LLM-based approach chosen for flexibility, context awareness, and natural language.",
         "student_answer": "LLMs are best because they understand context and can generate diverse questions."},
        {"score": 0, "justification": "No answer was provided by the student.",
         "correct_answer_hint": "Limitations include API dependency, inability to handle scanned PDFs, and token limits.",
         "student_answer": ""},
    ]

    # Generate test score sheet
    output = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "..", "outputs", "score_sheets", "test_scoresheet.docx"
    )
    output = os.path.normpath(output)

    result = generate_score_sheet(
        student_name="John Doe",
        student_id="2024CS001",
        subject="Machine Learning",
        department="Computer Science",
        report_title="Automated Viva Voce Question Generator Using LLMs and Bloom's Taxonomy",
        faculty_name="Dr. Jane Smith",
        questions_data=sample_questions,
        eval_results=sample_results,
        output_path=output,
    )

    print(f"\n  ✓ Score sheet generated at: {result}")
    print(f"  ✓ File exists: {os.path.isfile(result)}")
    print(f"  ✓ File size: {os.path.getsize(result)} bytes")

    print("\n" + "=" * 60)
    print("DOCX generator test passed ✓")
    print("=" * 60)
